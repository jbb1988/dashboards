#!/usr/bin/env python3
"""
Generate a revised DOCX by modifying the ORIGINAL document.
Uses python-docx (FREE, no watermark) to preserve formatting.

Input (JSON from stdin):
{
  "originalDocxBase64": "base64 encoded docx",
  "originalText": "original text for diffing",
  "modifiedText": "AI-modified text"
}

Output (JSON to stdout):
{
  "success": true,
  "docxBase64": "base64 encoded result",
  "changesApplied": 5
}
"""

import sys
import json
import base64
import re
from io import BytesIO

from docx import Document
from diff_match_patch import diff_match_patch


def normalize_for_matching(text: str) -> str:
    """Normalize text for matching."""
    # Normalize quotes
    text = re.sub(r'[\u201C\u201D\u201E\u201F\u2033\u2036]', '"', text)
    text = re.sub(r'[\u2018\u2019\u201A\u201B\u2032\u2035]', "'", text)
    # Normalize dashes
    text = re.sub(r'[\u2013\u2014\u2015]', '-', text)
    # Collapse whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def extract_replacements(original_text: str, modified_text: str) -> list:
    """Extract find/replace pairs from the diff."""
    dmp = diff_match_patch()

    norm_original = normalize_for_matching(original_text)
    norm_modified = normalize_for_matching(modified_text)

    diffs = dmp.diff_main(norm_original, norm_modified)
    dmp.diff_cleanupSemantic(diffs)

    replacements = []
    i = 0

    while i < len(diffs):
        op, text = diffs[i]

        if op == 0:  # EQUAL
            i += 1
            continue

        if op == -1:  # DELETE
            deleted = text
            inserted = ""

            # Check if followed by INSERT
            if i + 1 < len(diffs) and diffs[i + 1][0] == 1:
                inserted = diffs[i + 1][1]
                i += 1

            if deleted.strip() or inserted.strip():
                # Get context
                context_before = ""
                for j in range(i - 1, -1, -1):
                    if diffs[j][0] == 0:
                        ctx = diffs[j][1]
                        if len(ctx) > 40:
                            ctx = ctx[-40:]
                            space_pos = ctx.find(' ')
                            if space_pos > 0:
                                ctx = ctx[space_pos + 1:]
                        context_before = ctx
                        break

                context_after = ""
                for j in range(i + 1, len(diffs)):
                    if diffs[j][0] == 0:
                        ctx = diffs[j][1]
                        if len(ctx) > 40:
                            ctx = ctx[:40]
                            space_pos = ctx.rfind(' ')
                            if space_pos > 0:
                                ctx = ctx[:space_pos]
                        context_after = ctx
                        break

                find_text = context_before + deleted + context_after
                replace_text = context_before + inserted + context_after

                if find_text.strip() and find_text != replace_text:
                    replacements.append({
                        'find': find_text.strip(),
                        'replace': replace_text.strip(),
                        'deleted': deleted.strip()[:50],
                        'inserted': inserted.strip()[:50]
                    })

        i += 1

    return replacements


def replace_in_paragraph(paragraph, find_text: str, replace_text: str) -> bool:
    """Replace text in a paragraph while preserving formatting."""
    # Get full paragraph text
    full_text = paragraph.text

    # Normalize for matching
    norm_full = normalize_for_matching(full_text)
    norm_find = normalize_for_matching(find_text)

    if norm_find not in norm_full:
        return False

    # Find the position in normalized text
    pos = norm_full.find(norm_find)
    if pos == -1:
        return False

    # For simple cases, do direct replacement
    # This preserves the paragraph's runs but replaces text
    for run in paragraph.runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text)
            return True
        # Also try normalized matching
        norm_run = normalize_for_matching(run.text)
        if norm_find in norm_run:
            # Replace with smart quotes preserved
            run.text = run.text.replace(run.text, run.text.replace(
                next((s for s in [find_text, norm_find] if s in run.text), find_text),
                replace_text
            ))
            return True

    return False


def generate_revised_docx(original_base64: str, original_text: str, modified_text: str) -> dict:
    """Modify the original DOCX with AI changes, preserving formatting."""
    try:
        # Decode original document
        original_bytes = base64.b64decode(original_base64)

        if original_bytes[:4] != b'PK\x03\x04':
            return {'success': False, 'error': 'Invalid DOCX file'}

        # Load document with python-docx
        doc_stream = BytesIO(original_bytes)
        doc = Document(doc_stream)

        # Extract replacements from diff
        replacements = extract_replacements(original_text, modified_text)
        print(f"Found {len(replacements)} replacements", file=sys.stderr)

        applied_count = 0

        for rep in replacements:
            find_text = rep['find']
            replace_text = rep['replace']

            # Skip short matches
            if len(find_text) < 20:
                continue

            # Try to apply in each paragraph
            for para in doc.paragraphs:
                if replace_in_paragraph(para, find_text, replace_text):
                    applied_count += 1
                    print(f"Applied: '{rep['deleted'][:30]}' -> '{rep['inserted'][:30]}'", file=sys.stderr)
                    break

            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if replace_in_paragraph(para, find_text, replace_text):
                                applied_count += 1
                                break

        # Save to bytes
        output_stream = BytesIO()
        doc.save(output_stream)
        output_bytes = output_stream.getvalue()

        result_base64 = base64.b64encode(output_bytes).decode('utf-8')

        return {
            'success': True,
            'docxBase64': result_base64,
            'changesApplied': applied_count,
            'changesTotal': len(replacements)
        }

    except Exception as e:
        import traceback
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }


def main():
    try:
        input_data = json.load(sys.stdin)

        original_base64 = input_data.get('originalDocxBase64', '')
        original_text = input_data.get('originalText', '')
        modified_text = input_data.get('modifiedText', '')

        if not original_base64:
            print(json.dumps({'success': False, 'error': 'originalDocxBase64 required'}))
            return

        if not original_text or not modified_text:
            print(json.dumps({'success': False, 'error': 'originalText and modifiedText required'}))
            return

        result = generate_revised_docx(original_base64, original_text, modified_text)
        print(json.dumps(result))

    except json.JSONDecodeError as e:
        print(json.dumps({'success': False, 'error': f'Invalid JSON: {str(e)}'}))
    except Exception as e:
        print(json.dumps({'success': False, 'error': f'Error: {str(e)}'}))


if __name__ == '__main__':
    main()
